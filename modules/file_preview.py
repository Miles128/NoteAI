import base64
import os
from typing import Any

from utils.helpers import get_file_extension, read_file_with_encoding
from utils.logger import logger


class FilePreviewer:
    SUPPORTED_PREVIEW_TYPES = [".md", ".markdown", ".txt", ".pdf", ".doc", ".docx"]

    def __init__(self, workspace_path: str | None = None):
        self.workspace_path = workspace_path

    def can_preview(self, file_path: str) -> bool:
        ext = get_file_extension(file_path).lower()
        return ext in self.SUPPORTED_PREVIEW_TYPES

    def get_preview_data(self, file_path: str) -> dict[str, Any]:
        full_path = file_path
        if not os.path.isabs(full_path) and self.workspace_path:
            full_path = os.path.join(self.workspace_path, file_path)

        if not os.path.exists(full_path):
            return {"success": False, "error": "文件不存在"}

        ext = get_file_extension(full_path).lower()
        file_size = os.path.getsize(full_path)

        try:
            if ext == ".md" or ext == ".markdown":
                return self._preview_markdown(full_path, file_size)
            if ext == ".txt":
                return self._preview_text(full_path, file_size)
            if ext == ".pdf":
                return self._preview_pdf(full_path, file_size)
            if ext == ".docx":
                return self._preview_docx(full_path, file_size)
            if ext == ".doc":
                return self._preview_doc_legacy(full_path, file_size)
            return {"success": False, "error": f"不支持预览此文件类型: {ext}"}
        except Exception as e:
            logger.error(f"预览失败 {full_path}: {e}")
            return {"success": False, "error": str(e)}

    def _utf8_transport(self, content: str) -> dict[str, Any]:
        """Encode preview text once as standard base64 (ASCII-safe over JSON-RPC)."""
        raw = content.encode("utf-8")
        return {
            "transport": "base64_utf8",
            "content_b64": base64.standard_b64encode(raw).decode("ascii"),
            "content_byte_length": len(raw),
        }

    def _preview_markdown(self, file_path: str, file_size: int) -> dict[str, Any]:
        content = read_file_with_encoding(file_path)

        return {
            "success": True,
            "type": "markdown",
            "preview_delivery": "semantic_b64",
            "file_name": os.path.basename(file_path),
            "file_size": file_size,
            **self._utf8_transport(content),
        }

    def _preview_text(self, file_path: str, file_size: int) -> dict[str, Any]:
        content = read_file_with_encoding(file_path)

        return {
            "success": True,
            "type": "text",
            "preview_delivery": "semantic_b64",
            "file_name": os.path.basename(file_path),
            "file_size": file_size,
            **self._utf8_transport(content),
        }

    def _preview_pdf(self, file_path: str, file_size: int) -> dict[str, Any]:
        try:
            import fitz
        except ImportError:
            return self._preview_pdf_legacy(file_path, file_size)

        pages_data = []
        total_pages = 0
        full_text = ""
        MAX_PREVIEW_PAGES = 50
        MAX_PAGE_PIXELS = 3000 * 3000  # 单页最大像素限制，防止超大页面导致内存耗尽
        MAX_TOTAL_IMAGE_SIZE = 200 * 1024 * 1024  # 总图片数据上限 200MB
        total_image_size = 0

        doc = None
        try:
            doc = fitz.open(file_path)
            total_pages = len(doc)
            preview_pages = min(total_pages, MAX_PREVIEW_PAGES)

            for page_num in range(preview_pages):
                page = doc[page_num]

                text = page.get_text("text") or ""

                mat = fitz.Matrix(1.5, 1.5)
                # 检查页面像素尺寸，防止恶意超大 PDF 导致内存耗尽
                page_rect = page.rect
                page_width_px = int(page_rect.width * 1.5)
                page_height_px = int(page_rect.height * 1.5)
                if page_width_px * page_height_px > MAX_PAGE_PIXELS:
                    # 缩小渲染比例
                    scale = (MAX_PAGE_PIXELS / (page_rect.width * page_rect.height)) ** 0.5
                    mat = fitz.Matrix(scale, scale)

                pix = page.get_pixmap(matrix=mat, alpha=False)
                img_bytes = pix.tobytes("png")
                total_image_size += len(img_bytes)
                if total_image_size > MAX_TOTAL_IMAGE_SIZE:
                    break
                img_base64 = base64.b64encode(img_bytes).decode("utf-8")

                pages_data.append(
                    {
                        "page_number": page_num + 1,
                        "text": text,
                        "image": img_base64,
                        "width": pix.width,
                        "height": pix.height,
                    }
                )

                full_text += text + "\n\n"

            return {
                "success": True,
                "type": "pdf",
                "file_name": os.path.basename(file_path),
                "file_size": file_size,
                "total_pages": total_pages,
                "pages": pages_data,
                "full_text": full_text.strip(),
                "truncated": total_pages > MAX_PREVIEW_PAGES,
            }
        except Exception as e:
            logger.error(f"PDF预览失败 (PyMuPDF): {e}")
            return self._preview_pdf_legacy(file_path, file_size)
        finally:
            if doc is not None:
                doc.close()

    def _preview_pdf_legacy(self, file_path: str, file_size: int) -> dict[str, Any]:
        try:
            from utils.pdf_utils import extract_pdf_pages

            pages_text = extract_pdf_pages(file_path)
            pages_data = []
            for i, text in enumerate(pages_text):
                pages_data.append({"page_number": i + 1, "text": text.strip(), "image": None, "width": 0, "height": 0})
            return {
                "success": True,
                "type": "pdf",
                "file_name": os.path.basename(file_path),
                "file_size": file_size,
                "total_pages": len(pages_data),
                "pages": pages_data,
                "full_text": "\n\n".join([p["text"] for p in pages_data]),
            }
        except Exception as e:
            logger.error(f"PDF预览失败 (legacy): {e}")
            return {"success": False, "error": f"PDF解析失败: {str(e)}"}

    def _preview_docx(self, file_path: str, file_size: int) -> dict[str, Any]:
        """DOCX → HTML（mammoth），供前端排版预览。"""
        try:
            import mammoth
        except ImportError:
            return self._preview_docx_python_docx(file_path, file_size)

        try:
            with open(file_path, "rb") as docx_file:
                result = mammoth.convert_to_html(docx_file)
            html = (result.value or "").strip()
            if not html:
                return self._preview_docx_python_docx(file_path, file_size)

            payload: dict[str, Any] = {
                "success": True,
                "type": "docx",
                "content_kind": "html",
                "preview_delivery": "semantic_b64",
                "file_name": os.path.basename(file_path),
                "file_size": file_size,
                **self._utf8_transport(html),
            }
            messages = [str(m) for m in getattr(result, "messages", [])]
            if messages:
                payload["warnings"] = messages[:8]
            return payload
        except Exception as e:
            logger.warning(f"DOCX mammoth 预览失败，回退 python-docx: {e}")
            return self._preview_docx_python_docx(file_path, file_size)

    def _preview_docx_python_docx(self, file_path: str, file_size: int) -> dict[str, Any]:
        from docx import Document

        try:
            doc = Document(file_path)
            parts: list[str] = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                style_name = (para.style.name if para.style else "Normal") or "Normal"
                if style_name.startswith("Heading"):
                    level = "".join(c for c in style_name if c.isdigit()) or "1"
                    parts.append(f"<h{level}>{self._escape_html(text)}</h{level}>")
                else:
                    parts.append(f"<p>{self._escape_html(text)}</p>")

            for table in doc.tables:
                rows_html: list[str] = []
                for row in table.rows:
                    cells = "".join(f"<td>{self._escape_html(cell.text.strip())}</td>" for cell in row.cells)
                    rows_html.append(f"<tr>{cells}</tr>")
                if rows_html:
                    parts.append("<table>" + "".join(rows_html) + "</table>")

            html = "".join(parts) or "<p>（文档无可见正文）</p>"
            return {
                "success": True,
                "type": "docx",
                "content_kind": "html",
                "preview_delivery": "semantic_b64",
                "file_name": os.path.basename(file_path),
                "file_size": file_size,
                **self._utf8_transport(html),
            }
        except Exception as e:
            logger.error(f"DOCX 预览失败: {e}")
            return {
                "success": False,
                "error": f"Word 文档解析失败: {str(e)}",
            }

    def _preview_doc_legacy(self, file_path: str, file_size: int) -> dict[str, Any]:
        """旧版 .doc：提取为 Markdown 后预览。"""
        try:
            from modules.file_converter import LegacyDOCConverter

            markdown_content = LegacyDOCConverter().to_markdown(file_path)
            return {
                "success": True,
                "type": "docx",
                "content_kind": "markdown",
                "preview_delivery": "semantic_b64",
                "file_name": os.path.basename(file_path),
                "file_size": file_size,
                **self._utf8_transport(markdown_content or ""),
            }
        except Exception as e:
            logger.error(f"DOC 预览失败: {e}")
            return {
                "success": False,
                "error": f"Word 文档解析失败: {str(e)}",
            }

    @staticmethod
    def _escape_html(text: str) -> str:
        return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace('"', "&quot;")


def get_file_type_display_name(ext: str) -> str:
    type_map = {
        ".md": "Markdown",
        ".markdown": "Markdown",
        ".txt": "文本文件",
        ".pdf": "PDF文档",
        ".doc": "Word文档",
        ".docx": "Word文档",
    }
    return type_map.get(ext.lower(), "未知类型")
