"""
Integration-style checks for path resolution, topic tree payload, and preview path rules.

Requires project dependencies (see pyproject.toml). Run: pytest tests/integration/
"""

from __future__ import annotations

import base64
import json
import threading
from pathlib import Path
from types import SimpleNamespace

import pytest
from sidecar.handlers.files_handler import FilesHandler
from sidecar.handlers.tags_handler import TagsHandler
from sidecar.handlers.topics_handler import TopicsHandler
from sidecar.handlers.workspace_handler import WorkspaceHandler
from sidecar.paths import find_file_by_name_in_workspace, resolve_workspace_path
from sidecar.pending_topics import load_pending_topics
from sidecar.rag.index import _rag_index_dir
from sidecar.server import WATCHED_WORKSPACE_SUFFIXES, SidecarServer

from config import config
from config.constants import TOPIC_SEP
from config.settings import RAG_INDEX_FOLDER, WORKSPACE_APP_FOLDER
from modules.file_preview import FilePreviewer
from utils.activity_log import add_entry, get_entries
from utils.topic_assigner import (
    auto_assign_topic_for_file,
    parse_wiki_headings,
    parse_wiki_structure,
    sync_wiki_with_files,
)
from utils.topic_manager import TopicManager
from utils.wiki_manager import topic_from_notes_path


@pytest.fixture
def workspace(tmp_path: Path) -> Path:
    d = tmp_path / "ws"
    d.mkdir()
    (d / "Notes").mkdir()
    (d / "wiki").mkdir(parents=True, exist_ok=True)
    config.workspace_path = str(d)
    return d


class TestResolveWorkspacePath:
    def test_relative_path_inside_workspace(self, workspace: Path) -> None:
        rel = "Notes/hello.md"
        f = workspace / rel
        f.write_text("# x", encoding="utf-8")
        got = resolve_workspace_path(rel)
        assert got == str(f.resolve())

    def test_absolute_path_inside_workspace(self, workspace: Path) -> None:
        f = workspace / "Notes" / "a.md"
        f.write_text("y", encoding="utf-8")
        got = resolve_workspace_path(str(f))
        assert got == str(f.resolve())

    def test_rejects_escape_outside_workspace(self, workspace: Path) -> None:
        outside = workspace.parent / "evil.md"
        outside.write_text("z", encoding="utf-8")
        assert resolve_workspace_path(str(outside)) is None

    def test_returns_none_when_workspace_unset(self, monkeypatch: pytest.MonkeyPatch) -> None:
        monkeypatch.setattr(config, "workspace_path", "", raising=False)
        assert resolve_workspace_path("Notes/x.md") is None


class TestFindFileByName:
    def test_finds_first_match_in_workspace(self, workspace: Path) -> None:
        (workspace / "wiki" / "dup.md").write_text("1", encoding="utf-8")
        got = find_file_by_name_in_workspace("wiki/dup.md")
        assert got is not None
        assert got.endswith("dup.md")


class TestWorkspaceTree:
    def test_watched_workspace_suffixes_match_user_visible_inputs(self) -> None:
        assert {".md", ".txt", ".pdf", ".docx", ".pptx", ".html", ".doc", ".ppt"} == WATCHED_WORKSPACE_SUFFIXES

    def test_only_supported_file_types_are_returned(self, workspace: Path) -> None:
        notes = workspace / "Notes"
        for name in ("a.md", "b.txt", "c.pdf", "d.docx", "e.pptx", "f.html", "g.doc", "h.ppt"):
            (notes / name).write_text("x", encoding="utf-8")
        for name in ("ignore.png", "ignore.jpg", "ignore.json"):
            (notes / name).write_text("x", encoding="utf-8")
        (workspace / "empty").mkdir()
        hidden = workspace / ".hidden"
        hidden.mkdir()
        (hidden / "secret.md").write_text("x", encoding="utf-8")
        node_modules = workspace / "node_modules"
        node_modules.mkdir()
        (node_modules / "package-note.md").write_text("x", encoding="utf-8")
        noteai = workspace / WORKSPACE_APP_FOLDER
        noteai.mkdir()
        (noteai / "GUIDE.md").write_text("x", encoding="utf-8")
        (workspace / "wiki" / "WIKI.md").write_text("x", encoding="utf-8")

        handler = WorkspaceHandler(SimpleNamespace(_ctx=SimpleNamespace(config=config, logger=None)))
        tree = handler._compute_workspace_tree()
        notes_node = next(node for node in tree if node["name"] == "Notes")
        names = {node["name"] for node in notes_node["children"]}

        assert names == {"a.md", "b.txt", "c.pdf", "d.docx", "e.pptx", "f.html", "g.doc", "h.ppt"}
        root_names = {node["name"] for node in tree}
        assert "empty" not in root_names
        assert ".hidden" not in root_names
        assert "node_modules" not in root_names
        assert WORKSPACE_APP_FOLDER not in root_names
        assert "wiki" in root_names

    def test_workspace_watcher_treats_topic_directory_deletion_as_relevant(self, workspace: Path) -> None:
        server = SidecarServer.__new__(SidecarServer)

        assert server._is_relevant_workspace_change(workspace / "Notes" / "主题", is_directory=True)
        assert not server._is_relevant_workspace_change(workspace / ".trash" / "主题", is_directory=True)
        assert not server._is_relevant_workspace_change(workspace / "wiki" / "主题", is_directory=True)
        assert not server._is_relevant_workspace_change(workspace / WORKSPACE_APP_FOLDER, is_directory=True)

    def test_workspace_watcher_syncs_wiki_for_external_note_deletions(self, workspace: Path) -> None:
        server = SidecarServer.__new__(SidecarServer)

        note_path = workspace / "Notes" / "主题" / "文件.md"
        assert server._workspace_change_affects_wiki("deleted", note_path, is_directory=False)
        assert server._workspace_change_affects_wiki("deleted", workspace / "Notes" / "主题", is_directory=True)
        assert not server._workspace_change_affects_wiki("modified", note_path, is_directory=False)
        assert not server._workspace_change_affects_wiki("deleted", workspace / "wiki" / "WIKI.md", is_directory=False)
        assert not server._workspace_change_affects_wiki(
            "deleted", workspace / WORKSPACE_APP_FOLDER / "log.md", is_directory=False
        )

    def test_topic_from_notes_path_reads_folder_hierarchy(self, workspace: Path) -> None:
        note = workspace / "Notes" / "AI" / "产品" / "一篇.md"
        note.parent.mkdir(parents=True, exist_ok=True)
        note.write_text("x", encoding="utf-8")

        assert topic_from_notes_path(note) == f"AI{TOPIC_SEP}产品"
        assert topic_from_notes_path(workspace / "Notes" / "孤立.md") is None

    def test_watcher_keeps_file_when_moved_to_new_topic_folder(self, workspace: Path) -> None:
        server = SidecarServer.__new__(SidecarServer)
        server._watcher_debounce_lock = threading.Lock()
        server._watcher_needs_wiki_sync = False

        new_dir = workspace / "Notes" / "AI" / "新主题"
        new_dir.mkdir(parents=True, exist_ok=True)
        note = new_dir / "笔记.md"
        note.write_text(f"---\ntopic: AI{TOPIC_SEP}旧主题\n---\n正文\n", encoding="utf-8")

        server._auto_process_md_file(str(note))

        assert note.exists()
        text = note.read_text(encoding="utf-8")
        assert f"AI{TOPIC_SEP}新主题" in text
        assert f"AI{TOPIC_SEP}旧主题" not in text

    def test_sync_wiki_follows_notes_folder_after_move(self, workspace: Path) -> None:
        old = workspace / "Notes" / "主题A"
        new = workspace / "Notes" / "主题B"
        old.mkdir(parents=True)
        note = old / "文件.md"
        note.write_text("---\ntopic: 主题A\n---\n", encoding="utf-8")

        new.mkdir(parents=True, exist_ok=True)
        note.rename(new / "文件.md")

        sync_wiki_with_files()
        wiki = (workspace / "wiki" / "WIKI.md").read_text(encoding="utf-8")

        assert "主题B" in wiki
        assert "**文件**" in wiki
        assert "主题A" not in wiki or "1. **文件**" not in wiki.split("主题A")[0]

    def test_rag_index_lives_under_workspace_noteai_folder(self, workspace: Path) -> None:
        assert _rag_index_dir(str(workspace)) == workspace / WORKSPACE_APP_FOLDER / RAG_INDEX_FOLDER

        ok, _message = config.setup_workspace_folders()

        assert ok is True
        assert (workspace / WORKSPACE_APP_FOLDER / RAG_INDEX_FOLDER).is_dir()
        assert WORKSPACE_APP_FOLDER.startswith(".")

    def test_activity_log_writes_unified_wiki_log(self, workspace: Path) -> None:
        add_entry("test", "hello", "detail")

        log_path = workspace / "wiki" / "log.md"
        entries = get_entries(10)

        assert log_path.exists()
        assert entries[-1]["msg"] == "hello"
        assert "hello" in log_path.read_text(encoding="utf-8")


class TestTagsHandler:
    def test_auto_tag_files_only_updates_frontmatter(self, workspace: Path) -> None:
        notes = workspace / "Notes"
        source = notes / "source.md"
        source.write_text("---\ntags:\n- LangGraph\n---\nsource body", encoding="utf-8")
        target = notes / "LangGraph 实战.md"
        body = "\n# LangGraph 实战\n\n| A | B |\n|---|---|\n| 1 | 2 |\n"
        target.write_text(body, encoding="utf-8")
        handler = TagsHandler(SimpleNamespace(_ctx=SimpleNamespace(config=config, logger=None)))

        result = handler._auto_tag_files({"dry_run": False})

        assert result["updated"] == 1
        updated = target.read_text(encoding="utf-8")
        assert "tags:\n- LangGraph" in updated
        assert "# LangGraph 实战\n\n| A | B |\n|---|---|\n| 1 | 2 |" in updated

    def test_delete_tag_preserves_body(self, workspace: Path) -> None:
        note = workspace / "Notes" / "tagged.md"
        body = "\n# Title\n\nOriginal body\n"
        note.write_text("---\ntags:\n- A\n- B\n---" + body, encoding="utf-8")
        handler = TagsHandler(SimpleNamespace(_ctx=SimpleNamespace(config=config, logger=None)))

        result = handler._delete_tag({"tag_name": "A"})

        assert result["updated"] == 1
        updated = note.read_text(encoding="utf-8")
        assert "tags:\n- B" in updated
        assert "# Title\n\nOriginal body" in updated


class TestParseWikiStructure:
    """parse_wiki_structure() parses WIKI.md into a list of topic dicts."""

    def test_returns_topic_list(self, workspace: Path) -> None:
        wiki_dir = workspace / "wiki"
        wiki_dir.mkdir(exist_ok=True)
        wiki = wiki_dir / "WIKI.md"
        wiki.write_text(
            "## AI产品经理之路\n\n1. **产品思维**\n2. **需求分析**\n\n### Agent 架构\n\n1. **Agent 设计**\n",
            encoding="utf-8",
        )
        topics = parse_wiki_structure()
        assert isinstance(topics, list)
        min_topic_count = 2
        assert len(topics) >= min_topic_count
        for t in topics:
            assert "name" in t
            assert "label" in t
            assert "files" in t
            assert isinstance(t["files"], list)

    def test_empty_when_no_wiki(self, workspace: Path) -> None:
        _ = workspace
        topics = parse_wiki_structure()
        assert topics == []

    def test_wiki_sync_uses_notes_folder_as_source_of_truth(self, workspace: Path) -> None:
        topic_dir = workspace / "Notes" / "普通人的AI指南" / "Agent 入门"
        topic_dir.mkdir(parents=True)
        note = topic_dir / "提示词.md"
        note.write_text("---\ntopic: 旧主题\n---\n正文", encoding="utf-8")
        root_note = workspace / "Notes" / "未分类.md"
        root_note.write_text("---\ntopic: 不应保留\n---\n正文", encoding="utf-8")

        result = sync_wiki_with_files()

        assert result["success"] is True
        wiki = (workspace / "wiki" / "WIKI.md").read_text(encoding="utf-8")
        assert "## 普通人的AI指南" in wiki
        assert "### Agent 入门" in wiki
        assert "1. **提示词**" in wiki
        assert "未分类" not in wiki
        assert "topic: 普通人的AI指南 > Agent 入门" in note.read_text(encoding="utf-8")
        assert "topic:" not in root_note.read_text(encoding="utf-8")

    def test_parse_wiki_headings_returns_full_topic_paths(self, workspace: Path) -> None:
        (workspace / "wiki" / "WIKI.md").write_text(
            "## 普通人的AI指南\n\n### Agent 入门\n\n#### 工具调用\n",
            encoding="utf-8",
        )

        names = [item["name"] for item in parse_wiki_headings()]

        assert names == [
            "普通人的AI指南",
            "普通人的AI指南 > Agent 入门",
            "普通人的AI指南 > Agent 入门 > 工具调用",
        ]


class TestPendingTopics:
    """load_pending_topics() reads .pending_topics.json from workspace."""

    def test_pending_json_roundtrip(self, workspace: Path) -> None:
        pending_path = workspace / ".pending_topics.json"
        sample = [{"file": "Notes/a.md", "title": "A", "candidates": ["T1"]}]
        pending_path.write_text(json.dumps(sample, ensure_ascii=False), encoding="utf-8")
        pending = load_pending_topics()
        assert len(pending) == 1
        assert pending[0]["file"] == "Notes/a.md"

    def test_empty_when_no_file(self, workspace: Path) -> None:
        _ = workspace
        pending = load_pending_topics()
        assert pending == []


class TestPreviewPathContract:
    """
    Preview pipeline: resolve → absolute path → FilePreviewer.get_preview_data.
    Relative paths must still work when workspace_path is set on FilePreviewer.
    """

    def _decoded_semantic_preview(self, data: dict) -> str:
        assert data.get("transport") == "base64_utf8"
        return base64.standard_b64decode(data["content_b64"]).decode("utf-8")

    def test_absolute_resolved_path_matches_existing_file(self, workspace: Path) -> None:
        rel = "Notes/preview.md"
        f = workspace / rel
        f.write_text("# Title\n\nbody", encoding="utf-8")
        abs_path = resolve_workspace_path(rel)
        assert abs_path is not None
        prev = FilePreviewer(str(workspace))
        data = prev.get_preview_data(abs_path)
        assert data.get("success") is True
        assert data.get("type") == "markdown"
        assert "Title" in self._decoded_semantic_preview(data)

    def test_relative_path_joined_with_workspace(self, workspace: Path) -> None:
        rel = "Notes/rel.md"
        (workspace / rel).write_text("# R", encoding="utf-8")
        prev = FilePreviewer(str(workspace))
        data = prev.get_preview_data(rel)
        assert data.get("success") is True
        assert "# R" in self._decoded_semantic_preview(data)

    def test_missing_file_returns_error_not_success(self, workspace: Path) -> None:
        prev = FilePreviewer(str(workspace))
        data = prev.get_preview_data("Notes/nope.md")
        assert data.get("success") is False
        assert data.get("error") == "文件不存在"

    def test_large_markdown_preview_uses_binary_slices(self, workspace: Path) -> None:
        rel = "Notes/big.md"
        expected = "# H\n\n" + ("测" * 190_000)
        (workspace / rel).write_text(expected, encoding="utf-8")

        srv = SimpleNamespace(
            _ctx=SimpleNamespace(config=config, logger=None),
            file_previewer=FilePreviewer(str(workspace)),
            _resolve_path=lambda path: resolve_workspace_path(path),
            _find_file_by_name=lambda name: find_file_by_name_in_workspace(name),
        )
        handler = FilesHandler(srv)

        head = handler._get_file_preview({"path": rel})
        assert head.get("success") is True
        assert head.get("preview_delivery") == "raw_slices"

        blobs: list[bytes] = []
        offset = 0
        guard = 0
        while offset < head["total_byte_size"] and guard < 50:
            guard += 1
            slab = handler._read_preview_raw_slice({"path": rel, "byte_offset": offset, "byte_limit": 16384})
            assert slab["success"] is True
            blobs.append(base64.standard_b64decode(slab["chunk_b64"]))
            offset = slab["next_byte_offset"]
            if slab["done"]:
                break

        merged = b"".join(blobs).decode("utf-8")
        assert merged == expected

        semantic = handler._get_file_preview({"path": rel, "force_semantic_preview": True})
        assert semantic.get("preview_delivery") == "semantic_b64"
        assert self._decoded_semantic_preview(semantic) == expected


class TestWorkspaceTreeContract:
    """RPC-shaped workspace tree payload for the file sidebar."""

    def test_get_workspace_tree_includes_standard_roots(self, workspace: Path) -> None:
        (workspace / "Raw").mkdir(exist_ok=True)
        (workspace / "Notes" / "note.md").write_text("# n\n", encoding="utf-8")

        srv = SimpleNamespace(_ctx=SimpleNamespace(config=config, logger=None))
        handler = WorkspaceHandler(srv)
        tree = handler._get_workspace_tree({})

        assert isinstance(tree, list)
        names = {item["name"] for item in tree if item.get("type") == "folder"}
        assert "Notes" in names
        assert "wiki" in names
        assert "Raw" in names

    def test_get_workspace_tree_includes_docx_under_raw(self, workspace: Path) -> None:
        docx_path = workspace / "Raw" / "paper.docx"
        docx_path.parent.mkdir(parents=True, exist_ok=True)
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not installed")

        doc = Document()
        doc.add_paragraph("smoke")
        doc.save(str(docx_path))

        srv = SimpleNamespace(_ctx=SimpleNamespace(config=config, logger=None))
        handler = WorkspaceHandler(srv)
        tree = handler._get_workspace_tree({})

        raw = next(item for item in tree if item["name"] == "Raw")
        file_names = [c["name"] for c in raw.get("children", []) if c.get("type") == "file"]
        assert "paper.docx" in file_names


class TestFilesHandlerPreviewContract:
    """FilesHandler preview RPC contract for DOCX (frontend preview path)."""

    def _handler(self, workspace: Path) -> FilesHandler:
        return FilesHandler(
            SimpleNamespace(
                _ctx=SimpleNamespace(config=config, logger=None),
                file_previewer=FilePreviewer(str(workspace)),
                _resolve_path=lambda path: resolve_workspace_path(path),
                _find_file_by_name=lambda name: find_file_by_name_in_workspace(name),
            )
        )

    def test_can_preview_file_docx(self, workspace: Path) -> None:
        rel = "Raw/sample.docx"
        path = workspace / rel
        path.parent.mkdir(parents=True, exist_ok=True)
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not installed")

        doc = Document()
        doc.add_heading("Contract", level=1)
        doc.save(str(path))

        handler = self._handler(workspace)
        assert handler._can_preview_file({"path": rel}) is True

    def test_get_file_preview_docx_semantic_html(self, workspace: Path) -> None:
        rel = "Raw/preview.docx"
        path = workspace / rel
        path.parent.mkdir(parents=True, exist_ok=True)
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not installed")

        doc = Document()
        doc.add_paragraph("contract body")
        doc.save(str(path))

        handler = self._handler(workspace)
        data = handler._get_file_preview({"path": rel})

        assert data.get("success") is True
        assert data.get("type") == "docx"
        assert data.get("content_kind") == "html"
        html = base64.standard_b64decode(data["content_b64"]).decode("utf-8")
        assert "contract" in html.lower()


class TestTopicAssignFromNotesFolder:
    def test_auto_assign_derives_topic_from_notes_path(self, workspace: Path) -> None:
        topic_dir = workspace / "Notes" / "普通人的AI指南" / "二级测试"
        topic_dir.mkdir(parents=True)
        note = topic_dir / "笔记.md"
        note.write_text("# 标题\n正文\n", encoding="utf-8")

        result = auto_assign_topic_for_file(str(note), use_llm=False)

        assert result is not None
        assert result.get("status") == "auto_assigned"
        assert result.get("topic") == "普通人的AI指南 > 二级测试"
        assert result.get("source") == "folder_path"


class TestCollectTopicLabelsForPendingUi:
    def test_collect_topic_labels_three_levels(self, workspace: Path) -> None:
        l1 = "普通人的AI指南"
        (workspace / "Notes" / l1 / "二级A" / "三级1").mkdir(parents=True)
        labels = TopicManager.collect_topic_labels(str(workspace))
        assert l1 in labels
        assert f"{l1}{TOPIC_SEP}二级A" in labels
        assert f"{l1}{TOPIC_SEP}二级A{TOPIC_SEP}三级1" in labels

    def test_get_all_pending_includes_topic_options(self, workspace: Path) -> None:
        l1 = "普通人的AI指南"
        (workspace / "Notes" / l1 / "待选二级").mkdir(parents=True)

        srv = SimpleNamespace(_ctx=SimpleNamespace(config=config, logger=None))
        handler = TopicsHandler(srv)
        payload = handler._get_all_pending({})

        assert "topic_options" in payload
        opts = payload["topic_options"]
        assert isinstance(opts, list)
        assert f"{l1}{TOPIC_SEP}待选二级" in opts


class TestIngestAndSchemaHandlers:
    def test_ensure_schema_and_get_schema(self, workspace: Path) -> None:
        from sidecar.handlers.ingest_handler import IngestHandler

        srv = SimpleNamespace(
            _ctx=SimpleNamespace(config=config, logger=None),
            _running_tasks=set(),
            _running_tasks_lock=__import__("threading").Lock(),
        )
        handler = IngestHandler(srv)

        assert handler._needs_schema_setup({})["needs_setup"] is True
        created = handler._ensure_schema({})
        assert created["success"] is True

        handler._save_schema({"content": "# test schema\nai_may_edit_wiki: true\n"})
        assert (workspace / "schema.md").exists()
        assert handler._needs_schema_setup({})["needs_setup"] is False

        got = handler._get_schema({})
        assert got["success"] is True
        assert "ai_may_edit_wiki" in got["content"]

    def test_get_ingest_status_idle(self, workspace: Path) -> None:
        from sidecar.handlers.ingest_handler import IngestHandler

        srv = SimpleNamespace(_ctx=SimpleNamespace(config=config, logger=None))
        handler = IngestHandler(srv)
        status = handler._get_ingest_status({})
        assert status["success"] is True
        assert status["status"] in ("idle", "complete", "failed", "cancelled", "running")
