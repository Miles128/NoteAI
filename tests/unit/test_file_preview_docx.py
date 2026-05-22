"""DOCX preview returns HTML payload for the web UI."""

from __future__ import annotations

import base64
from pathlib import Path

import pytest

from modules.file_preview import FilePreviewer


def _decoded(data: dict) -> str:
    assert data.get("transport") == "base64_utf8"
    return base64.standard_b64decode(data["content_b64"]).decode("utf-8")


@pytest.fixture
def sample_docx(tmp_path: Path) -> Path:
    docx = tmp_path / "sample.docx"
    try:
        from docx import Document
    except ImportError:
        pytest.skip("python-docx not installed")

    doc = Document()
    doc.add_heading("预览标题", level=1)
    doc.add_paragraph("第一段正文。")
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "A1"
    table.rows[0].cells[1].text = "B1"
    doc.save(str(docx))
    return docx


def test_docx_preview_returns_html(sample_docx: Path) -> None:
    prev = FilePreviewer(str(sample_docx.parent))
    data = prev.get_preview_data(str(sample_docx))

    assert data.get("success") is True
    assert data.get("type") == "docx"
    assert data.get("content_kind") == "html"
    html = _decoded(data)
    assert "预览标题" in html or "第一段" in html


def test_can_preview_docx(sample_docx: Path) -> None:
    prev = FilePreviewer()
    assert prev.can_preview(str(sample_docx)) is True
